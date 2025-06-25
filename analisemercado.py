#!/usr/bin/env python
# coding: utf-8

import os
import re
import uuid
import logging
import markdown
import matplotlib.pyplot as plt
from datetime import datetime
from dotenv import load_dotenv
from crewai import Agent, Task, Crew
from crewai_tools import SerperDevTool,FileWriterTool,ScrapeWebsiteTool
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE

# Configuração de logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Carregar variáveis de ambiente
load_dotenv()

# Ferramentas
serper = SerperDevTool()
scraper = ScrapeWebsiteTool()
file_writer = FileWriterTool()

# Configuração do relatório
CONFIG_RELATORIO = {
    "paginas_minimas": 10,
    "contagem_palavras": 10000,
    "secoes": [
        "Descrição da Categoria", "Panorama de Mercado", "Análise da Cadeia de Suprimentos",
        "Indicadores Econômicos", "Benchmarking", "Análise SWOT", "5 Forças de Porter", "CBD (Cost Breakdown)",
        "SCA (Should Cost Analysis)", "LPP (Line Performance Pricing)", "Perguntas para Clientes e Fornecedores",
        "SLAs e Multas", "Critérios de Seleção de Fornecedores", "Estrutura de RFP", "Análise de Riscos",
        "Alavancas de Negociação e BATNA", "Momento Ideal para Negociação", "Tendências ESG", "Conclusão"
    ]
}

# Função para criar gráfico aranha
def criar_grafico_aranha(dados, rotulos, titulo, arquivo):
    angulos = [n / float(len(rotulos)) * 2 * 3.14159 for n in range(len(rotulos))]
    angulos += angulos[:1]
    valores = list(dados.values()) + list(dados.values())[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angulos, valores, color='blue', alpha=0.25)
    ax.set_xticks(angulos[:-1])
    ax.set_xticklabels(rotulos)
    ax.set_title(titulo)
    plt.savefig(arquivo, format='png')
    plt.close()
    logging.info(f"Gráfico salvo como {arquivo}")

# Função auxiliar para adicionar tabela Markdown no Word
def adicionar_tabela_md_no_word(linhas_md, doc):
    linhas = [linha for linha in linhas_md if not re.match(r'^\|[-\s|]+\|$', linha)]
    if not linhas:
        return

    header = [cell.strip() for cell in linhas[0].split('|') if cell.strip()]
    tabela = doc.add_table(rows=1, cols=len(header))
    tabela.autofit = True
    hdr_cells = tabela.rows[0].cells
    for i, texto in enumerate(header):
        hdr_cells[i].text = texto

    for linha in linhas[1:]:
        valores = [cell.strip() for cell in linha.split('|') if cell.strip()]
        row_cells = tabela.add_row().cells
        for i, texto in enumerate(valores):
            if i < len(row_cells):
                row_cells[i].text = texto

# Função para exportar conteúdo Markdown para documento Word
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import os

def exportar_para_word(conteudo_md, setor, arquivo_saida):
    try:
        doc = Document()
        doc.add_heading(f'Relatório de Análise de Mercado: {setor}', 0)
        doc.add_paragraph('Preparado por Vorätte Consultoria')

        if 'TituloNegrito1' not in doc.styles:
            estilo_titulo1 = doc.styles.add_style('TituloNegrito1', WD_STYLE_TYPE.PARAGRAPH)
            estilo_titulo1.font.bold = True
            estilo_titulo1.font.size = Pt(14)

        if 'TituloNegrito2' not in doc.styles:
            estilo_titulo2 = doc.styles.add_style('TituloNegrito2', WD_STYLE_TYPE.PARAGRAPH)
            estilo_titulo2.font.bold = True
            estilo_titulo2.font.size = Pt(12)

        linhas = conteudo_md.split('\n')
        i = 0
        while i < len(linhas):
            linha = linhas[i].strip()
            if not linha or linha == '---':
                i += 1
                continue

            # Ignorar imagens markdown ![...](...)
            if re.match(r'^!\[.*\]\(.*\)', linha):
                i += 1
                continue

            # Ignorar linhas de markdown bruto ou blocos de código
            if linha.startswith('`') or linha.lower().strip() == 'markdown':
                i += 1
                continue

            if linha.startswith('|') and '|' in linha:
                tabela_linhas = [linha]
                i += 1
                while i < len(linhas) and linhas[i].strip().startswith('|'):
                    tabela_linhas.append(linhas[i].strip())
                    i += 1
                adicionar_tabela_md_no_word(tabela_linhas, doc)
                continue

            if linha.startswith('# '):
                doc.add_paragraph(linha[2:].strip(), style='TituloNegrito1')
            elif linha.startswith('## '):
                doc.add_paragraph(linha[3:].strip(), style='TituloNegrito2')
            elif linha.startswith('### '):
                doc.add_paragraph(linha[4:].strip(), style='TituloNegrito2')
            else:
                linha_limpa = re.sub(r'\*\*(.*?)\*\*', r'\1', linha)
                linha_limpa = re.sub(r'\*(.*?)\*', r'\1', linha_limpa)
                linha_limpa = re.sub(r'^[-\*\+]\s+', '', linha_limpa)
                linha_limpa = re.sub(r'[–—]+', '', linha_limpa)
                linha_limpa = re.sub(r'`{1,3}', '', linha_limpa)
                linha_limpa = re.sub(r'\(.*?\)', '', linha_limpa)  # remove conteúdo entre parênteses (de links)
                linha_limpa = re.sub(r'\[.*?\]', '', linha_limpa)  # remove conteúdo entre colchetes (de links)
                doc.add_paragraph(linha_limpa.strip())

            i += 1

        if os.path.exists('swot_aranha.png'):
            doc.add_picture('swot_aranha.png', width=Inches(4))
        if os.path.exists('porter_aranha.png'):
            doc.add_picture('porter_aranha.png', width=Inches(4))

        doc.save(arquivo_saida)
        print(f"Documento Word salvo como: {arquivo_saida}")

    except Exception as e:
        print(f"Erro ao exportar para Word: {e}")
        raise

def adicionar_tabela_md_no_word(linhas_md, doc):
    linhas = [linha for linha in linhas_md if not re.match(r'^\|[-\s|]+\|$', linha)]
    if not linhas:
        return

    header = [cell.strip() for cell in linhas[0].split('|') if cell.strip()]
    tabela = doc.add_table(rows=1, cols=len(header))
    tabela.autofit = True
    hdr_cells = tabela.rows[0].cells
    for i, texto in enumerate(header):
        hdr_cells[i].text = texto

    for linha in linhas[1:]:
        valores = [cell.strip() for cell in linha.split('|') if cell.strip()]
        row_cells = tabela.add_row().cells
        for i, texto in enumerate(valores):
            if i < len(row_cells):
                row_cells[i].text = texto

# Agentes
pesquisador = Agent(
    role="Pesquisador de Mercado",
    goal="Coletar dados completos sobre {setor} em {regiao} para {empresa}, cobrindo oferta/demanda, competidores, custos, forças, fortalezas, indicadores econômicos, ESG e riscos.",
    backstory="Analista experiente em sourcing estratégico, especializado em dados de mercado.",
    tools=[serper, scraper],
    verbose=True
)

analista = Agent(
    role="Analista de Tendências",
    goal="Analisar dados de {setor} para gerar insights, incluindo SWOT, 5 Forças de Porter, TCO, LPP, CBD, SCA e tendências ESG.",
    backstory="Especialista em estratégias de compras, negociação, focado em insights acionáveis.",
    tools=[serper],
    verbose=True
)

redator = Agent(
    role="Redator de Relatórios",
    goal="Criar um relatório detalhado em Markdown sobre {setor}, com visualizações, exportado para Word.",
    backstory="Redator profissional que transforma dados complexos em relatórios estratégicos.",
    verbose=True
)

# Tarefas
coleta_dados = Task(
    description=(
        "Coletar dados sobre {setor} em {regiao} para {empresa}, incluindo:\n"
        "1. Panorama de oferta e demanda (global e regional).\n"
        "2. Top 20 competidores por faturamento (2024 ou mais recente).\n"
        "3. Modelos de contratação (ex.: pay-as-you-go, reservado).\n"
        "4. Estrutura de custos e tendências de preços (últimos 24 meses, vs. USD, IGPM, IPCA).\n"
        "5. CBD, SCA, TCO, LPP, SLAs e multas.\n"
        "6. Detalhes da cadeia de suprimentos (produtores, logística, dependência de importação).\n"
        "7. Práticas e tendências ESG.\n"
        "8. Riscos (geopolíticos, regulatórios, etc.) e oportunidades.\n"
        "9. Estimar dados ausentes com base em tendências recentes."
    ),
    expected_output="JSON estruturado com dados brutos para todos os elementos solicitados.",
    agent=pesquisador
)

analise_tendencias = Task(
    description=(
        "Analisar dados de {setor} para produzir:\n"
        "1. Análise SWOT com gráfico aranha.\n"
        "2. 5 Forças de Porter com gráfico aranha.\n"
        "3. Modelo TCO e fórmula LPP.\n"
        "4. Avaliação de riscos e estratégias de mitigação.\n"
        "5. Alavancas de negociação e BATNA.\n"
        "6. Momento ideal para negociação com base em sazonalidade."
    ),
    expected_output="JSON com insights, incluindo nomes dos arquivos de gráficos.",
    agent=analista
)

redacao_relatorio = Task(
    description=(
        f"Gerar um relatório de {CONFIG_RELATORIO['contagem_palavras']} palavras em Markdown para {{setor}}, "
        f"cobrindo todas as seções: {', '.join(CONFIG_RELATORIO['secoes'])}.\n"
        "Incluir tabelas, gráficos aranha e recomendações acionáveis.\n"
        "Exportar para Word com identidade visual da Vorätte.\n"
        f"Garantir que o relatório seja em português, detalhado, com no mínimo {CONFIG_RELATORIO['paginas_minimas']} páginas.\n"
        "Basear-se em dados atuais, com projeções para 12-18 meses."
    ),
    expected_output="Arquivo Markdown e documento Word (.docx).",
    agent=redator
)

# Equipe
equipe = Crew(
    agents=[pesquisador, analista, redator],
    tasks=[coleta_dados, analise_tendencias, redacao_relatorio],
    verbose=True
)

# Função principal
def gerar_relatorio(setor, regiao, empresa):
    try:
        resultado = equipe.kickoff(inputs={
            "setor": setor,
            "regiao": regiao,
            "empresa": empresa
        })

        # Gerar gráficos de exemplo
        swot_dados = {"Escalabilidade": 9, "Custo": 6, "Inovação": 8, "Risco": 5}
        porter_dados = {
            "Poder dos Fornecedores": 5,
            "Poder dos Compradores": 8,
            "Novos Entrantes": 3,
            "Substitutos": 6,
            "Rivalidade": 9
        }

        criar_grafico_aranha(swot_dados, swot_dados.keys(), "Análise SWOT", "swot_aranha.png")
        criar_grafico_aranha(porter_dados, porter_dados.keys(), "5 Forças de Porter", "porter_aranha.png")

        arquivo_saida = f"relatorio_{setor.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        exportar_para_word(str(resultado), setor, arquivo_saida)

        return arquivo_saida
    except Exception as e:
        logging.error(f"Erro ao gerar relatório: {e}")
        raise

# Execução
if __name__ == "__main__":
    resultado = gerar_relatorio("Serviços de Computação em Nuvem", "Brasil", "Grande Empresa")
    print(resultado.raw)
